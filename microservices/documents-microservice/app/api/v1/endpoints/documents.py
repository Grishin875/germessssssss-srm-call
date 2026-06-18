import logging, os, uuid, time, random, string
from typing import Optional
from datetime import datetime
from fastapi import APIRouter, Request, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy import select, update, delete, func, text, or_
from pydantic import BaseModel

from app.models.documents import FirmwareBatch, Document

logger = logging.getLogger(__name__)

STORAGE_DIR = "/app/storage/documents"

MIME_MAP = {
    "pdf":  "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "jpg":  "image/jpeg",
    "jpeg": "image/jpeg",
    "png":  "image/png",
}


import re as _re


def _html_to_runs(paragraph, html: str):
    """Parse simple HTML (b, i, u, br) into python-docx runs."""
    from docx.shared import Pt
    # Replace <br> with newline marker
    html = _re.sub(r'<br\s*/?>', '\n', html)
    # Tokenise: split on bold/italic/underline tags
    pattern = _re.compile(r'(</?(?:b|strong|i|em|u)>)', _re.IGNORECASE)
    parts = pattern.split(html)
    bold = italic = underline = False
    for part in parts:
        low = part.lower()
        if low in ('<b>', '<strong>'):   bold = True;      continue
        if low in ('</b>', '</strong>'): bold = False;     continue
        if low in ('<i>', '<em>'):       italic = True;    continue
        if low in ('</i>', '</em>'):     italic = False;   continue
        if low == '<u>':                 underline = True;  continue
        if low == '</u>':                underline = False; continue
        # strip remaining tags
        text = _re.sub(r'<[^>]+>', '', part)
        if not text:
            continue
        for chunk in text.split('\n'):
            run = paragraph.add_run(chunk)
            run.bold = bold
            run.italic = italic
            run.underline = underline


def _html_to_plain(html: str) -> str:
    """Strip HTML tags, preserve line breaks."""
    html = _re.sub(r'<br\s*/?>', '\n', html, flags=_re.IGNORECASE)
    html = _re.sub(r'</p>', '\n', html, flags=_re.IGNORECASE)
    html = _re.sub(r'<[^>]+>', '', html)
    return html.strip()


def _esc(text: str) -> str:
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def _docx_para_to_html(p, img_parts: dict, img_budget: list) -> str:
    """Convert a python-docx Paragraph to an HTML string."""
    from docx.oxml.ns import qn
    import base64

    style_name = (p.style.name or '').lower()
    tag = 'p'
    if 'heading 1' in style_name:   tag = 'h1'
    elif 'heading 2' in style_name: tag = 'h2'
    elif 'heading 3' in style_name: tag = 'h3'
    elif 'heading 4' in style_name: tag = 'h4'

    parts = []
    for child in p._element:
        child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if child_tag == 'r':  # run
            run_text = ''.join(n.text or '' for n in child.iter() if n.tag.split('}')[-1] == 't')
            text = _esc(run_text)
            # check formatting
            rpr = child.find(qn('w:rPr'))
            if rpr is not None:
                if rpr.find(qn('w:b')) is not None:   text = f'<b>{text}</b>'
                if rpr.find(qn('w:i')) is not None:   text = f'<i>{text}</i>'
                if rpr.find(qn('w:u')) is not None:   text = f'<u>{text}</u>'
                color_el = rpr.find(qn('w:color'))
                if color_el is not None:
                    color_val = color_el.get(qn('w:val'), '')
                    if color_val and color_val.lower() != 'auto':
                        text = f'<span style="color:#{color_val}">{text}</span>'
            parts.append(text)
        elif child_tag == 'drawing':  # inline image
            for blip in child.iter(qn('a:blip')):
                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId and rId in img_parts and img_budget[0] < 3_000_000:
                    rel = img_parts[rId]
                    try:
                        data = base64.b64encode(rel.blob).decode()
                        mime = rel.content_type or 'image/png'
                        img_budget[0] += len(data)
                        parts.append(f'<img src="data:{mime};base64,{data}" style="max-width:100%;height:auto;display:block;margin:8px 0">')
                    except Exception:
                        logger.warning("Не удалось извлечь изображение из docx", exc_info=True)

    inner = ''.join(parts).strip()
    if not inner:
        return '<br>'
    return f'<{tag} style="margin:6px 0;line-height:1.6">{inner}</{tag}>'


def _docx_table_to_html(tbl) -> str:
    from docx.table import Table
    rows_html = []
    for row in tbl.rows:
        cells_html = []
        for cell in row.cells:
            cell_text = _esc(cell.text)
            cells_html.append(f'<td style="padding:5px 10px;border:1px solid #cbd5e1;vertical-align:top">{cell_text}</td>')
        rows_html.append(f'<tr>{"".join(cells_html)}</tr>')
    return f'<table style="border-collapse:collapse;width:100%;font-size:13px;margin:12px 0">{"".join(rows_html)}</table>'


def _extract_content(file_path: str, file_type: str) -> str:
    """Extract content as HTML (preserves formatting, tables, and inline images)."""
    try:
        if file_type == "docx":
            from docx import Document as DocxDoc
            from docx.oxml.ns import qn
            doc = DocxDoc(file_path)

            # Build a map of relationship id → image part
            img_parts = {}
            try:
                for rId, rel in doc.part.rels.items():
                    if 'image' in rel.reltype:
                        img_parts[rId] = rel._target
            except Exception:
                logger.warning("Не удалось построить карту изображений docx", exc_info=True)

            img_budget = [0]  # mutable counter for total base64 image bytes
            blocks = []
            para_map = {p._element: p for p in doc.paragraphs}
            tbl_map  = {t._element: t for t in doc.tables}

            for child in doc.element.body:
                local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if local == 'p' and child in para_map:
                    blocks.append(_docx_para_to_html(para_map[child], img_parts, img_budget))
                elif local == 'tbl' and child in tbl_map:
                    blocks.append(_docx_table_to_html(tbl_map[child]))

            return '\n'.join(blocks)

        elif file_type == "pdf":
            import fitz
            with fitz.open(file_path) as pdf:
                pages = []
                for page in pdf:
                    text = page.get_text()
                    escaped = _esc(text)
                    pages.append(f'<div style="margin-bottom:24px;white-space:pre-wrap;font-family:monospace;font-size:13px">{escaped}</div>')
                return '\n'.join(pages)

        elif file_type == "xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True)
            ws = wb.active
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = '\t'.join(str(c) if c is not None else '' for c in row)
                rows.append(cells)
            return '<br>'.join(rows)

    except Exception:
        logger.exception("Не удалось извлечь содержимое файла %s (%s)", file_path, file_type)
    return ''


def _save_content_to_file(file_path: str, file_type: str, content: str):
    """Save HTML content back into the file with formatting."""
    if file_type == "docx":
        from docx import Document as DocxDoc
        from docx.shared import Pt, RGBColor
        doc = DocxDoc()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        # Split on <br> / <p> tags into paragraphs
        blocks = _re.split(r'<br\s*/?>\s*|</p>', content, flags=_re.IGNORECASE)
        for block in blocks:
            block = block.strip()
            if not block:
                doc.add_paragraph('')
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _html_to_runs(p, block)
        doc.save(file_path)

    elif file_type == "pdf":
        import os
        import fpdf as _fpdf_module
        from fpdf import FPDF
        # Use DejaVu font (ships with fpdf2) for Cyrillic support
        fonts_dir = os.path.join(os.path.dirname(_fpdf_module.__file__), 'fonts')
        pdf = FPDF()
        pdf.set_margins(20, 20, 20)
        pdf.add_page()
        pdf.add_font('DejaVu', fname=os.path.join(fonts_dir, 'DejaVuSans.ttf'))
        pdf.add_font('DejaVu', style='B', fname=os.path.join(fonts_dir, 'DejaVuSans-Bold.ttf'))
        pdf.add_font('DejaVu', style='I', fname=os.path.join(fonts_dir, 'DejaVuSans-Oblique.ttf'))
        pdf.add_font('DejaVu', style='BI', fname=os.path.join(fonts_dir, 'DejaVuSans-BoldOblique.ttf'))
        pdf.set_font('DejaVu', size=12)
        plain = _html_to_plain(content)
        for line in plain.split('\n'):
            pdf.multi_cell(0, 8, line if line else ' ')
        pdf.output(file_path)

    elif file_type == "xlsx":
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        plain = _html_to_plain(content)
        for row_str in plain.split('\n'):
            ws.append(row_str.split('\t'))
        wb.save(file_path)

router = APIRouter()


def _db(r): return r.state.db
def _user(r):
    u = r.state.current_user
    if not u: raise HTTPException(401, "Не авторизован")
    return u

def _op_id(p): return f"{p}-{int(time.time()*1000)}-{''.join(random.choices(string.ascii_lowercase,k=6))}"

def _m(obj) -> dict:
    return {c.key: getattr(obj, c.key) for c in obj.__mapper__.column_attrs}


class FirmwareCreate(BaseModel):
    source_batch_id: str
    product_name: str
    qty: int
    operator_id: str
    firmware_version: Optional[str] = None
    comment: Optional[str] = None


class FirmwareComplete(BaseModel):
    good_qty: int
    defect_qty: int
    comment: Optional[str] = None


class FirmwareUpdate(BaseModel):
    firmware_version: Optional[str] = None
    status: Optional[str] = None
    comment: Optional[str] = None


# ── Firmware ──────────────────────────────────────────────────────────────────

def _fw_condition(fw_id: str):
    """fw_id может быть числовым id или строковым batch_id — Integer-колонку
    нельзя сравнивать с нечисловой строкой."""
    try:
        fw_int_id = int(fw_id)
    except (ValueError, TypeError):
        return FirmwareBatch.batch_id == fw_id
    return or_(FirmwareBatch.id == fw_int_id, FirmwareBatch.batch_id == fw_id)


@router.get("/firmware")
async def list_firmware(request: Request, status: Optional[str] = None,
                        source_batch_id: Optional[str] = None):
    _user(request)
    db = _db(request)
    q = select(FirmwareBatch)
    if status:
        q = q.where(FirmwareBatch.status == status)
    if source_batch_id:
        q = q.where(FirmwareBatch.source_batch_id == source_batch_id)
    q = q.order_by(FirmwareBatch.start_date.desc(), FirmwareBatch.created_at.desc())
    try:
        result = await db.execute(q)
        batches = result.scalars().all()
        rows = []
        for b in batches:
            d = _m(b)
            if b.operator_id:
                op_name = (await db.execute(text(
                    "SELECT name FROM operators WHERE employee_id=:e"
                ), {"e": b.operator_id})).scalar_one_or_none()
                d["operator_name"] = op_name
                src = (await db.execute(text(
                    "SELECT product_name FROM production_batches WHERE batch_id=:b"
                ), {"b": b.source_batch_id})).scalar_one_or_none() if b.source_batch_id else None
                d["source_product"] = src
            else:
                d["operator_name"] = None
                d["source_product"] = None
            rows.append(d)
        return rows
    except Exception:
        logger.exception("Запрос списка не выполнен")
        return []


@router.get("/firmware/{fw_id}")
async def get_firmware(fw_id: str, request: Request):
    _user(request)
    db = _db(request)
    result = await db.execute(select(FirmwareBatch).where(_fw_condition(fw_id)))
    fw = result.scalar_one_or_none()
    if not fw: raise HTTPException(404, "Партия прошивки не найдена")
    d = _m(fw)
    if fw.operator_id:
        op_name = (await db.execute(text(
            "SELECT name FROM operators WHERE employee_id=:e"
        ), {"e": fw.operator_id})).scalar_one_or_none()
        d["operator_name"] = op_name
    return d


@router.post("/firmware", status_code=201)
async def create_firmware(body: FirmwareCreate, request: Request):
    _user(request)
    db = _db(request)
    src = (await db.execute(text(
        "SELECT * FROM production_batches WHERE batch_id=:b"
    ), {"b": body.source_batch_id})).mappings().one_or_none()
    if not src: raise HTTPException(400, "Исходная партия SMD не найдена")
    if src["production_type"] != "SMD": raise HTTPException(400, "Исходная партия должна быть типа SMD")
    if src["status"] != "Завершена": raise HTTPException(400, "Исходная партия должна быть завершена")
    # Учитываем платы, уже забранные другими прошивочными партиями из этого источника
    used = int((await db.execute(
        select(func.coalesce(func.sum(FirmwareBatch.qty), 0))
        .where(FirmwareBatch.source_batch_id == body.source_batch_id)
    )).scalar_one())
    available = int(src["actual_qty"] or 0) - used
    if body.qty > available:
        raise HTTPException(
            400, f"Недостаточно плат: в партии {src['actual_qty']}, уже в прошивке {used}, "
                 f"доступно {max(0, available)}, требуется {body.qty}")

    date_part = datetime.utcnow().strftime("%y%m%d")
    op_short = (body.operator_id or "000")[-3:].zfill(3)
    base = f"FW{date_part}-{op_short}"
    bid, suffix = base, 1
    while (await db.execute(
        select(FirmwareBatch.batch_id).where(FirmwareBatch.batch_id == bid)
    )).scalar_one_or_none():
        bid = f"{base}-{suffix}"; suffix += 1

    fw = FirmwareBatch(
        batch_id=bid, source_batch_id=body.source_batch_id,
        product_name=body.product_name.strip(), qty=body.qty,
        operator_id=body.operator_id, firmware_version=body.firmware_version,
        status="В работе", start_date=datetime.utcnow(), comment=body.comment,
    )
    db.add(fw)
    await db.flush()
    await db.refresh(fw)
    await db.commit()
    return _m(fw)


@router.post("/firmware/{fw_id}/complete")
async def complete_firmware(fw_id: str, body: FirmwareComplete, request: Request):
    _user(request)
    db = _db(request)
    result = await db.execute(select(FirmwareBatch).where(_fw_condition(fw_id)))
    fw = result.scalar_one_or_none()
    if not fw: raise HTTPException(404, "Партия прошивки не найдена")
    if fw.status != "В работе": raise HTTPException(400, "Партия уже завершена")
    if body.good_qty + body.defect_qty != int(fw.qty):
        raise HTTPException(400, f"Сумма годных ({body.good_qty}) и брака ({body.defect_qty}) должна равняться {fw.qty}")
    await db.execute(
        update(FirmwareBatch)
        .where(FirmwareBatch.id == fw.id)
        .values(good_qty=body.good_qty, defect_qty=body.defect_qty,
                status="Завершена", end_date=func.now(),
                comment=body.comment, updated_at=func.now())
    )
    await db.execute(text("""
        INSERT INTO operations (operation_type, component_name, quantity, note, operator_id, operation_id)
        VALUES ('FIRMWARE_COMPLETE', :pn, :qty, :note, :op, :oid)
    """), {"pn": fw.product_name, "qty": body.good_qty,
           "note": f"Брак: {body.defect_qty}, Партия: {fw.batch_id}",
           "op": fw.operator_id, "oid": _op_id("FW")})
    await db.commit()
    return {"success": True, "batch_id": fw.batch_id}


@router.put("/firmware/{fw_id}")
async def update_firmware(fw_id: str, body: FirmwareUpdate, request: Request):
    _user(request)
    db = _db(request)
    update_data = body.model_dump(exclude_none=True)
    if not update_data: raise HTTPException(400, "Нет данных для обновления")
    stmt = (
        update(FirmwareBatch)
        .where(_fw_condition(fw_id))
        .values(**update_data, updated_at=func.now())
        .returning(FirmwareBatch)
    )
    row = (await db.execute(stmt)).mappings().one_or_none()
    if not row: raise HTTPException(404, "Партия прошивки не найдена")
    await db.commit()
    return dict(row)


# ── SMD ───────────────────────────────────────────────────────────────────────

@router.get("/smd")
async def list_smd(request: Request, status: Optional[str] = None):
    _user(request)
    db = _db(request)
    where, params = ["pb.production_type='SMD'"], {}
    if status: where.append("pb.status=:st"); params["st"] = status
    try:
        rows = (await db.execute(text(f"""
            SELECT pb.*, o.name as operator_name
            FROM production_batches pb
            LEFT JOIN operators o ON pb.operator_id=o.employee_id
            WHERE {' AND '.join(where)}
            ORDER BY pb.start_date DESC
        """), params)).mappings().all()
        return list(rows)
    except Exception:
        logger.exception("Запрос списка не выполнен")
        return []


@router.get("/smd/{batch_id}")
async def get_smd_batch(batch_id: str, request: Request):
    _user(request)
    db = _db(request)
    row = (await db.execute(text("""
        SELECT pb.*, o.name as operator_name
        FROM production_batches pb
        LEFT JOIN operators o ON pb.operator_id=o.employee_id
        WHERE pb.batch_id=:b AND pb.production_type='SMD'
    """), {"b": batch_id})).mappings().one_or_none()
    if not row: raise HTTPException(404, "SMD партия не найдена")
    return dict(row)


# ── Documents ──────────────────────────────────────────────────────────────────

class DocumentUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    category: Optional[str] = None
    tags: Optional[str] = None


class DocumentContentUpdate(BaseModel):
    content: str


@router.get("/documents")
async def list_documents(request: Request,
                         search: Optional[str] = None,
                         category: Optional[str] = None,
                         file_type: Optional[str] = None):
    _user(request)
    db = _db(request)
    q = select(Document).order_by(Document.created_at.desc())
    if search:
        q = q.where(or_(Document.name.ilike(f"%{search}%"),
                        Document.description.ilike(f"%{search}%"),
                        Document.tags.ilike(f"%{search}%")))
    if category:
        q = q.where(Document.category == category)
    if file_type:
        q = q.where(Document.file_type == file_type)
    result = await db.execute(q)
    return [_m(d) for d in result.scalars().all()]


@router.post("/documents/upload", status_code=201)
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    name: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    category: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),
):
    user = _user(request)
    db = _db(request)
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in MIME_MAP:
        raise HTTPException(400, f"Неподдерживаемый формат: .{ext}")

    file_id = uuid.uuid4().hex
    file_path = os.path.join(STORAGE_DIR, f"{file_id}.{ext}")
    data = await file.read()
    with open(file_path, "wb") as f:
        f.write(data)

    content = _extract_content(file_path, ext)
    doc = Document(
        name=(name or file.filename or "Без названия").strip(),
        description=description,
        category=category,
        tags=tags,
        file_path=file_path,
        file_name=file.filename,
        file_type=ext,
        file_size=len(data),
        content=content,
        created_by=getattr(user, "id", None),
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)
    await db.commit()
    return _m(doc)


@router.get("/documents/{doc_id}")
async def get_document(doc_id: int, request: Request):
    _user(request)
    db = _db(request)
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc: raise HTTPException(404, "Документ не найден")
    return _m(doc)


@router.get("/documents/{doc_id}/download")
async def download_document(doc_id: int, request: Request):
    _user(request)
    db = _db(request)
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc: raise HTTPException(404, "Документ не найден")
    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(404, "Файл не найден на сервере")
    return FileResponse(
        doc.file_path,
        filename=doc.file_name or f"document.{doc.file_type}",
        media_type=MIME_MAP.get(doc.file_type, "application/octet-stream"),
    )


@router.put("/documents/{doc_id}")
async def update_document(doc_id: int, body: DocumentUpdate, request: Request):
    _user(request)
    db = _db(request)
    data = {k: v for k, v in body.model_dump().items() if v is not None}
    if not data: raise HTTPException(400, "Нет данных для обновления")
    data["updated_at"] = datetime.utcnow()
    stmt = update(Document).where(Document.id == doc_id).values(**data).returning(Document)
    row = (await db.execute(stmt)).mappings().one_or_none()
    if not row: raise HTTPException(404, "Документ не найден")
    await db.commit()
    return dict(row)


@router.put("/documents/{doc_id}/content")
async def update_document_content(doc_id: int, body: DocumentContentUpdate, request: Request):
    _user(request)
    db = _db(request)
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc: raise HTTPException(404, "Документ не найден")
    if doc.file_type in ("jpg", "jpeg", "png"):
        raise HTTPException(400, "Изображения нельзя редактировать как текст")
    if doc.file_path and doc.file_type in ("docx", "pdf", "xlsx"):
        try:
            _save_content_to_file(doc.file_path, doc.file_type, body.content)
        except Exception:
            # DB update proceeds regardless
            logger.exception("Не удалось сохранить содержимое в файл %s", doc.file_path)
    await db.execute(
        update(Document).where(Document.id == doc_id)
        .values(content=body.content, updated_at=datetime.utcnow())
    )
    await db.commit()
    return {"success": True}


@router.post("/documents/{doc_id}/reextract")
async def reextract_document_content(doc_id: int, request: Request):
    """Re-extract content from the original file (e.g. after extraction logic is improved)."""
    _user(request)
    db = _db(request)
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc: raise HTTPException(404, "Документ не найден")
    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(404, "Файл не найден на сервере")
    content = _extract_content(doc.file_path, doc.file_type)
    await db.execute(
        update(Document).where(Document.id == doc_id)
        .values(content=content, updated_at=datetime.utcnow())
    )
    await db.commit()
    return {"success": True, "content_length": len(content)}


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: int, request: Request):
    _user(request)
    db = _db(request)
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc: raise HTTPException(404, "Документ не найден")
    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    await db.execute(delete(Document).where(Document.id == doc_id))
    await db.commit()
    return {"success": True}


@router.get("/documents/{doc_id}/convert")
async def convert_document(doc_id: int, to: str, request: Request):
    _user(request)
    if to not in ("docx", "pdf", "xlsx"):
        raise HTTPException(400, f"Неподдерживаемый формат конвертации: {to}. Доступны: docx, pdf, xlsx")
    db = _db(request)
    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc: raise HTTPException(404, "Документ не найден")
    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(404, "Файл не найден")

    content = doc.content or _extract_content(doc.file_path, doc.file_type)
    out_path = os.path.join(STORAGE_DIR, f"convert_{uuid.uuid4().hex}.{to}")
    try:
        _save_content_to_file(out_path, to, content)
    except Exception:
        logger.exception("Конвертация документа %s в %s не удалась", doc_id, to)
        raise HTTPException(500, "Не удалось сконвертировать документ")

    from starlette.background import BackgroundTask
    base_name = (doc.file_name or "document").rsplit(".", 1)[0]
    return FileResponse(
        out_path,
        filename=f"{base_name}.{to}",
        media_type=MIME_MAP.get(to, "application/octet-stream"),
        background=BackgroundTask(os.remove, out_path),  # temp-файл удаляется после отдачи
    )
